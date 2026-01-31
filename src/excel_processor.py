"""
Excel Processing Module
Handles loading, processing, and merging Excel files from SurveyHeart
"""

from pathlib import Path
from typing import List, Dict, Tuple, Optional
import openpyxl
from openpyxl.styles import Font, Alignment
from openpyxl.utils import get_column_letter
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from PIL import Image, ImageDraw, ImageFont

from src.validators import clean_name, clean_email, parse_score, validate_row_data
from src.color_config import get_fill_for_test, TEST_COLORS

logger = logging.getLogger(__name__)

class ExcelProcessor:
    """Process and consolidate test results from multiple Excel files"""
    
    REQUIRED_COLUMNS = ['Full Name', 'Email', 'Score', 'Result', '%']
    
    def __init__(self, input_dir: str, output_dir: str):
        """
        Initialize the Excel processor
        
        Args:
            input_dir (str): Directory containing input XLSX files
            output_dir (str): Directory for output files
        """
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.test_data = {}  # {test_num: {email: {name, score}}}
        
    def find_column_index(self, sheet, column_names: List[str]) -> Optional[int]:
        """
        Find the index of a column by any of the possible names
        
        Args:
            sheet: openpyxl worksheet
            column_names (List[str]): Possible column names to search for
            
        Returns:
            int: Column index (1-based) or None if not found
        """
        for row in sheet.iter_rows(min_row=1, max_row=1):
            for cell_idx, cell in enumerate(row, 1):
                if cell.value and any(name.lower() in str(cell.value).lower() for name in column_names):
                    return cell_idx
        return None
    
    def load_test_file(self, filepath: Path, test_number: int) -> bool:
        """
        Load a single test file and extract data
        
        Args:
            filepath (Path): Path to the test XLSX file
            test_number (int): Test number (1-5)
            
        Returns:
            bool: True if successful
        """
        try:
            logger.info(f"Loading test {test_number} from {filepath.name}")
            wb = openpyxl.load_workbook(filepath, data_only=True)
            ws = wb.active
            
            # Find column indices
            name_col = self.find_column_index(ws, ['Full Name', 'Name', 'Participant'])
            email_col = self.find_column_index(ws, ['Email', 'E-mail', 'Email Address'])
            
            # For score, prioritize the column that matches this specific test
            # e.g., for Test 2, look for "Test 2 Score" first
            score_col = self.find_column_index(ws, [f'Test {test_number} Score', f'Test {test_number} Result'])
            
            # If not found, fall back to generic score columns
            if not score_col:
                score_col = self.find_column_index(ws, ['Score', 'Result', '%', 'Percentage'])
            
            if not all([name_col, email_col, score_col]):
                logger.error(f"Could not find required columns in {filepath.name}")
                logger.error(f"  Name col: {name_col}, Email col: {email_col}, Score col: {score_col}")
                return False
            
            logger.info(f"Columns found - Name: {name_col}, Email: {email_col}, Score: {score_col}")
            
            # Extract data
            self.test_data[test_number] = {}
            row_count = 0
            
            for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                full_name = clean_name(row[name_col - 1] if name_col <= len(row) else "")
                email = clean_email(row[email_col - 1] if email_col <= len(row) else "")
                score = parse_score(row[score_col - 1] if score_col <= len(row) else None)
                
                is_valid, error_msg = validate_row_data(full_name, email, score)
                
                if is_valid:
                    self.test_data[test_number][email] = {
                        'name': full_name,
                        'score': score
                    }
                    row_count += 1
                    # Log first few records to verify correct file
                    if row_count <= 3:
                        logger.info(f"  Test {test_number} row {row_idx}: {full_name} = {score}")
                else:
                    logger.warning(f"Row {row_idx} in test {test_number}: {error_msg}")
            
            logger.info(f"Loaded {row_count} valid records from test {test_number}")
            wb.close()
            return True
            
        except Exception as e:
            logger.error(f"Error loading {filepath.name}: {str(e)}")
            return False
    
    def load_all_tests(self, max_tests: Optional[int] = None) -> int:
        """
        Load all test files from input directory dynamically
        
        Args:
            max_tests (int): Maximum number of tests to load. If None, loads all found tests.
            
        Returns:
            int: Number of tests successfully loaded
        """
        loaded_count = 0
        
        # Find all XLSX files and extract test numbers
        all_xlsx_files = sorted(self.input_dir.glob("*.xlsx"))
        test_nums = set()
        
        logger.info(f"Scanning {len(all_xlsx_files)} files in {self.input_dir}")
        
        for f in all_xlsx_files:
            # Extract test number from filename using flexible pattern
            test_num = self._extract_test_number_from_file(f.name)
            if test_num:
                test_nums.add(test_num)
                logger.debug(f"Found test {test_num} in file: {f.name}")
        
        logger.info(f"Found test numbers: {sorted(test_nums)}")
        
        if not test_nums:
            logger.warning("No test files found in directory")
            return 0
        
        # Log all files found in directory
        logger.info(f"All XLSX files in directory:")
        for f in all_xlsx_files:
            extracted_num = self._extract_test_number_from_file(f.name)
            logger.info(f"  {f.name} -> Test {extracted_num}")
        
        # Load ONLY the tests that were actually sent (not fill gaps)
        for test_num in sorted(test_nums):
            matching_file = self._find_test_file(test_num)
            
            if matching_file:
                logger.info(f"Loading test {test_num} from: {matching_file.name}")
                success = self.load_test_file(matching_file, test_num)
                if success:
                    loaded_count += 1
                    participant_count = len(self.test_data.get(test_num, {}))
                    logger.info(f"Successfully loaded test {test_num}: {participant_count} participants")
                    # Log participant emails for debugging
                    emails = list(self.test_data[test_num].keys())
                    logger.info(f"  Test {test_num} emails: {emails[:5]}{'...' if len(emails) > 5 else ''}")
                else:
                    logger.error(f"Failed to load test {test_num} from {matching_file.name}")
            else:
                logger.warning(f"Test {test_num} was detected but file not found (should not happen)")
        
        return loaded_count
    
    def _find_test_file(self, test_num: int) -> Optional[Path]:
        """Find the file matching a specific test number"""
        for f in sorted(self.input_dir.glob("*.xlsx")):
            if self._extract_test_number_from_file(f.name) == test_num:
                return f
        return None
    
    @staticmethod
    def _extract_test_number_from_file(filename: str) -> Optional[int]:
        """
        Extract test number from filename (matches _extract_test_number in telegram_bot)
        Supports: 'Test 1', 'test1', '1.xlsx', 'result_1', 'exam(1)', etc.
        """
        import re
        name_without_ext = filename.rsplit('.', 1)[0] if '.' in filename else filename
        
        # Try 1: Look for "Test N" or "test N" format first
        match = re.search(r'[Tt]est\s*(\d+)', name_without_ext)
        if match:
            return int(match.group(1))
        
        # Try 2: Look for any number in the filename
        match = re.search(r'(\d+)', name_without_ext)
        if match:
            return int(match.group(1))
        
        return None
        return loaded_count
    
    def validate_data_integrity(self) -> Dict:
        """
        Validate data integrity across all tests and return error report
        
        Returns:
            Dict with keys: 'valid', 'errors', 'warnings', 'missing_participants', 'name_mismatches', 'duplicate_scores'
        """
        report = {
            'valid': True,
            'errors': [],
            'warnings': [],
            'missing_participants': [],  # Email in base test but missing in other tests
            'name_mismatches': [],  # Same email with different names
            'duplicate_scores': [],  # Same email with identical scores across tests
        }
        
        if not self.test_data:
            report['valid'] = False
            report['errors'].append("No test data loaded")
            return report
        
        # Use FIRST available test as base (not hardcoded Test 1)
        available_tests = sorted(self.test_data.keys())
        base_test = available_tests[0]
        
        base_test_emails = set(self.test_data[base_test].keys())
        test_nums = sorted(self.test_data.keys())
        
        # Check each base test participant
        for email, base_test_data in self.test_data[base_test].items():
            base_test_name = base_test_data['name']
            base_test_score = base_test_data['score']
            
            # Check for missing in other tests
            for test_num in test_nums:
                if test_num != base_test:
                    if email not in self.test_data[test_num]:
                        report['missing_participants'].append({
                            'email': email,
                            'name': base_test_name,
                            'missing_in_test': test_num
                        })
            
            # Check for name mismatches across tests
            name_variants = {base_test_name}
            for test_num in test_nums:
                if test_num != base_test and email in self.test_data[test_num]:
                    other_name = self.test_data[test_num][email]['name']
                    if other_name.lower() != base_test_name.lower():
                        name_variants.add(other_name)
                        report['name_mismatches'].append({
                            'email': email,
                            'test_1_name': base_test_name,
                            'test_num': test_num,
                            'conflicting_name': other_name
                        })
            
            # Check for duplicate scores (possible copy-paste error)
            scores_by_test = {}
            scores_by_test[base_test] = base_test_score
            for test_num in test_nums:
                if test_num != base_test and email in self.test_data[test_num]:
                    scores_by_test[test_num] = self.test_data[test_num][email]['score']
            
            # Flag if all scores are identical
            unique_scores = set(scores_by_test.values())
            if len(unique_scores) == 1 and len(scores_by_test) > 1:
                report['duplicate_scores'].append({
                    'email': email,
                    'name': base_test_name,
                    'score': base_test_score,
                    'in_tests': sorted(scores_by_test.keys()),
                    'note': 'Identical scores in all tests - possible copy-paste error?'
                })
        
        # Set validity
        if report['errors'] or report['name_mismatches']:
            report['valid'] = False
        
        if report['missing_participants'] or report['duplicate_scores']:
            report['warnings'].append(
                f"{len(report['missing_participants'])} missing participant(s) and "
                f"{len(report['duplicate_scores'])} potential duplicate score(s)"
            )
        
        logger.info(f"Data integrity check: {len(report['missing_participants'])} missing, "
                   f"{len(report['name_mismatches'])} name mismatches, "
                   f"{len(report['duplicate_scores'])} duplicate scores")
        
        return report
    
    def consolidate_results(self) -> Dict:
        """
        Consolidate results from all tests into a single dataset dynamically
        
        Returns:
            Dict: Consolidated data {email: {name, test_N_score, ...}}
        """
        if not self.test_data:
            logger.warning("No test data loaded")
            return {}
        
        # Use the FIRST available test as base (not hardcoded Test 1)
        available_tests = sorted(self.test_data.keys())
        if not available_tests:
            logger.error("No test data available for consolidation")
            return {}
        
        base_test = available_tests[0]
        logger.info(f"=== CONSOLIDATION STARTING ===")
        logger.info(f"Using Test {base_test} as base for participant list")
        logger.info(f"Starting consolidation with {len(self.test_data)} test datasets")
        
        # Log ALL participants in ALL tests BEFORE consolidation
        logger.info("PRE-CONSOLIDATION PARTICIPANT BREAKDOWN:")
        for test_num in sorted(self.test_data.keys()):
            participants = list(self.test_data[test_num].keys())
            logger.info(f"  Test {test_num}: {len(participants)} participants")
            if len(participants) > 0:
                logger.info(f"    Sample: {participants[:3]}{'...' if len(participants) > 3 else ''}")
            else:
                logger.warning(f"  Test {test_num} has NO participants!")
        
        consolidated = {}
        
        # Iterate through base test participants (primary source)
        logger.info(f"Processing base test {base_test} participants...")
        base_participants = list(self.test_data[base_test].keys())
        logger.info(f"Base test {base_test} has {len(base_participants)} participants: {base_participants}")
        
        for email, data in self.test_data[base_test].items():
            consolidated[email] = {
                'name': data['name'],
                f'test_{base_test}_score': data['score']
            }
            
            # Add scores from all other tests dynamically by matching email
            for test_num in sorted(self.test_data.keys()):
                if test_num != base_test:
                    # Only add if email matches in that test
                    if email in self.test_data[test_num]:
                        score = self.test_data[test_num][email]['score']
                        consolidated[email][f'test_{test_num}_score'] = score
                    else:
                        # Explicitly set as None if not found
                        consolidated[email][f'test_{test_num}_score'] = None
        
        logger.info(f"After base test: {len(consolidated)} participants")
        
        # Sort by name
        consolidated = dict(sorted(consolidated.items(), 
                                  key=lambda x: x[1]['name'].lower()))
        
        logger.info(f"=== CONSOLIDATION COMPLETE ===")
        logger.info(f"Final consolidated: {len(consolidated)} participants across {len(self.test_data)} tests")
        if consolidated:
            first_email = list(consolidated.keys())[0]
            first_data = consolidated[first_email]
            logger.info(f"First participant ({first_email}): {first_data}")
        return consolidated
    
    def generate_preview_image(self, consolidated_data: Dict, max_rows: int = 12) -> Optional[Path]:
        """
        Generate a visual preview image showing consolidation summary and data table
        
        Args:
            consolidated_data (Dict): Consolidated results
            max_rows (int): Maximum rows to show in preview
            
        Returns:
            Path: Path to generated image file, or None if failed
        """
        try:
            # Get test numbers and stats
            test_nums = set()
            for data in consolidated_data.values():
                for key in data.keys():
                    if key.startswith('test_') and key.endswith('_score'):
                        test_nums.add(int(key.split('_')[1]))
            test_nums = sorted(test_nums)
            
            total_participants = len(consolidated_data)
            rows_to_show = min(max_rows, total_participants)
            
            # Image dimensions
            col_width = 150
            row_height = 30
            header_height = 100
            
            # Calculate image size
            num_cols = 2 + len(test_nums)  # Name + Email + Test scores
            img_width = col_width * num_cols + 20
            img_height = header_height + (rows_to_show + 1) * row_height + 80
            
            # Create image
            img = Image.new('RGB', (img_width, img_height), color='white')
            draw = ImageDraw.Draw(img)
            
            # Use default font (PIL will use a basic font)
            try:
                title_font = ImageFont.truetype("arial.ttf", 16)
                header_font = ImageFont.truetype("arial.ttf", 12)
                data_font = ImageFont.truetype("arial.ttf", 10)
            except:
                # Fallback to default font
                title_font = ImageFont.load_default()
                header_font = ImageFont.load_default()
                data_font = ImageFont.load_default()
            
            # Draw header
            draw.text((10, 5), "📊 CONSOLIDATION PREVIEW", fill='black', font=title_font)
            draw.text((10, 30), f"Participants: {total_participants} | Tests: {', '.join(f'T{t}' for t in test_nums)}", 
                     fill='#333333', font=data_font)
            
            # Draw column headers
            y = header_height
            x = 10
            headers = ['Name', 'Email'] + [f'Test {t}' for t in test_nums]
            
            for col_idx, header in enumerate(headers):
                # Alternate header background
                if col_idx % 2 == 0:
                    draw.rectangle([(x, y), (x + col_width, y + row_height)], fill='#E8E8E8')
                else:
                    draw.rectangle([(x, y), (x + col_width, y + row_height)], fill='#F5F5F5')
                
                draw.rectangle([(x, y), (x + col_width, y + row_height)], outline='#CCCCCC')
                draw.text((x + 5, y + 8), header[:15], fill='black', font=header_font)
                x += col_width
            
            y += row_height
            
            # Draw data rows
            for row_idx, (email, data) in enumerate(consolidated_data.items()):
                if row_idx >= rows_to_show:
                    break
                
                x = 10
                name = data['name'][:20]  # Truncate long names
                
                # Alternate row colors
                row_color = '#FFFFFF' if row_idx % 2 == 0 else '#F9F9F9'
                
                # Draw Name
                draw.rectangle([(x, y), (x + col_width, y + row_height)], fill=row_color)
                draw.rectangle([(x, y), (x + col_width, y + row_height)], outline='#DDDDDD')
                draw.text((x + 5, y + 8), name, fill='black', font=data_font)
                x += col_width
                
                # Draw Email
                draw.rectangle([(x, y), (x + col_width, y + row_height)], fill=row_color)
                draw.rectangle([(x, y), (x + col_width, y + row_height)], outline='#DDDDDD')
                email_display = email[:18] + '..' if len(email) > 18 else email
                draw.text((x + 5, y + 8), email_display, fill='#666666', font=data_font)
                x += col_width
                
                # Draw scores with color coding
                for test_num in test_nums:
                    score = data.get(f'test_{test_num}_score')
                    score_text = str(int(score)) if score is not None else '—'
                    
                    # Get test color
                    if test_num in TEST_COLORS:
                        color_hex = TEST_COLORS[test_num]['rgb']
                        # Convert hex to RGB
                        color_rgb = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
                    else:
                        color_rgb = (200, 200, 200)
                    
                    # Lighten color for background
                    bg_color = tuple(min(255, c + 100) for c in color_rgb)
                    
                    draw.rectangle([(x, y), (x + col_width, y + row_height)], fill=bg_color)
                    draw.rectangle([(x, y), (x + col_width, y + row_height)], outline='#DDDDDD')
                    draw.text((x + col_width//2 - 15, y + 8), score_text, fill='black', font=data_font)
                    x += col_width
                
                y += row_height
            
            # Draw footer
            if total_participants > rows_to_show:
                draw.text((10, y + 10), f"... and {total_participants - rows_to_show} more participants", 
                         fill='#666666', font=data_font)
            
            draw.text((10, y + 40), "✅ Scroll down to see validation alerts and confirm", 
                     fill='#228B22', font=data_font)
            
            # Save image
            output_path = Path(self.output_dir) / 'preview.png'
            img.save(output_path)
            logger.info(f"Generated preview image: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating preview image: {str(e)}")
            return None
    
    def save_consolidated_file(self, consolidated_data: Dict, output_filename: str = "Consolidated_Results.xlsx") -> bool:
        """
        Save consolidated results to Excel file with color coding (dynamic columns)
        
        Args:
            consolidated_data (Dict): Consolidated results
            output_filename (str): Output filename
            
        Returns:
            bool: True if successful
        """
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Results"
            
            # Get number of tests from data keys
            test_nums = []
            if consolidated_data:
                first_record = next(iter(consolidated_data.values()))
                for key in first_record:
                    if key.startswith('test_') and key.endswith('_score'):
                        test_num = int(key.split('_')[1])
                        test_nums.append(test_num)
                test_nums = sorted(test_nums)
            
            # Create headers dynamically
            headers = ['Full Name', 'Email'] + [f'Test {num} Score' for num in test_nums]
            ws.append(headers)
            
            # Format header row
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx)
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Add data rows
            for email, data in consolidated_data.items():
                row = [data['name'], email] + [data.get(f'test_{num}_score') for num in test_nums]
                ws.append(row)
            
            # Apply color coding to test score columns
            for row_idx in range(2, len(consolidated_data) + 2):
                for col_offset, test_num in enumerate(test_nums):
                    col_idx = col_offset + 3  # Column C onwards (A=name, B=email)
                    cell = ws.cell(row=row_idx, column=col_idx)
                    cell.fill = get_fill_for_test(test_num)
                    cell.alignment = Alignment(horizontal='center')
            
            # Adjust column widths
            ws.column_dimensions['A'].width = 25
            ws.column_dimensions['B'].width = 30
            for col_offset in range(len(test_nums)):
                col_letter = get_column_letter(col_offset + 3)
                ws.column_dimensions[col_letter].width = 15
            
            output_path = self.output_dir / output_filename
            wb.save(output_path)
            logger.info(f"Saved consolidated results to {output_path} ({len(test_nums)} tests)")
            return True
            
        except Exception as e:
            logger.error(f"Error saving consolidated file: {str(e)}")
            return False
    
    def save_as_pdf(self, consolidated_data: Dict, output_filename: str = "Consolidated_Results.pdf") -> bool:
        """
        Save consolidated results to PDF file (dynamic columns)
        
        Args:
            consolidated_data (Dict): Consolidated results
            output_filename (str): Output filename
            
        Returns:
            bool: True if successful
        """
        try:
            logger.info(f"Saving results as PDF: {output_filename}")
            
            # Get test numbers
            test_nums = []
            if consolidated_data:
                first_record = next(iter(consolidated_data.values()))
                for key in first_record:
                    if key.startswith('test_') and key.endswith('_score'):
                        test_num = int(key.split('_')[1])
                        test_nums.append(test_num)
                test_nums = sorted(test_nums)
            
            # Prepare data for table
            data = [['Full Name', 'Email'] + [f'Test {num}' for num in test_nums]]
            
            for email, record in consolidated_data.items():
                row = [
                    record['name'],
                    email
                ] + [str(record.get(f'test_{num}_score') or '') for num in test_nums]
                data.append(row)
            
            # Create PDF
            output_path = self.output_dir / output_filename
            doc = SimpleDocTemplate(str(output_path), pagesize=letter, topMargin=0.5*inch)
            
            # Calculate column widths dynamically
            col_widths = [1.8*inch, 2.2*inch] + [0.9*inch] * len(test_nums)
            table = Table(data, colWidths=col_widths)
            
            # Style table
            style_commands = [
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
            ]
            
            # Add color backgrounds to test score columns dynamically
            for row_idx in range(1, len(data)):
                for col_offset, test_num in enumerate(test_nums):
                    col_idx = col_offset + 2  # Starting from column C
                    color_hex = TEST_COLORS.get(test_num, {}).get('rgb', 'FFFFFF')
                    rgb = colors.HexColor(f'#{color_hex}')
                    style_commands.append(('BACKGROUND', (col_idx, row_idx), (col_idx, row_idx), rgb))
            
            table.setStyle(TableStyle(style_commands))
            
            # Build PDF
            elements = [table]
            doc.build(elements)
            
            logger.info(f"Saved PDF results to {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error saving PDF file: {str(e)}")
            return False
    
    def save_as_docx(self, consolidated_data: Dict, output_filename: str = "Consolidated_Results.docx") -> bool:
        """
        Save consolidated results to DOCX file (dynamic columns)
        
        Args:
            consolidated_data (Dict): Consolidated results
            output_filename (str): Output filename
            
        Returns:
            bool: True if successful
        """
        try:
            logger.info(f"Saving results as DOCX: {output_filename}")
            
            # Get test numbers
            test_nums = []
            if consolidated_data:
                first_record = next(iter(consolidated_data.values()))
                for key in first_record:
                    if key.startswith('test_') and key.endswith('_score'):
                        test_num = int(key.split('_')[1])
                        test_nums.append(test_num)
                test_nums = sorted(test_nums)
            
            # Create document
            doc = Document()
            doc.add_heading('Consolidated Test Results', 0)
            
            # Add summary
            doc.add_paragraph(f'Total Participants: {len(consolidated_data)}')
            doc.add_paragraph(f'Tests Included: {len(test_nums)}')
            doc.add_paragraph('')
            
            # Create table with dynamic columns
            num_cols = len(test_nums) + 2  # Name + Email + Tests
            table = doc.add_table(rows=1, cols=num_cols)
            table.style = 'Light Grid Accent 1'
            
            # Add header row
            header_cells = table.rows[0].cells
            headers = ['Full Name', 'Email'] + [f'Test {num}' for num in test_nums]
            
            for idx, header_text in enumerate(headers):
                header_cells[idx].text = header_text
                # Style header
                paragraph = header_cells[idx].paragraphs[0]
                paragraph.runs[0].font.bold = True
            
            # Add data rows
            for email, record in consolidated_data.items():
                row_cells = table.add_row().cells
                row_cells[0].text = record['name']
                row_cells[1].text = email
                
                for col_offset, test_num in enumerate(test_nums):
                    row_cells[col_offset + 2].text = str(record.get(f'test_{test_num}_score') or '')
                
                # Center align score columns
                for col_idx in range(2, 2 + len(test_nums)):
                    for paragraph in row_cells[col_idx].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            output_path = self.output_dir / output_filename
            doc.save(output_path)
            
            logger.info(f"Saved DOCX results to {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error saving DOCX file: {str(e)}")
            return False
