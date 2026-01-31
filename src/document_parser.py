#!/usr/bin/env python3
"""
Universal Document Parser for MLJ Results Compiler
Handles images, PDFs, Word docs, Excel, CSV, JSON, etc.
"""

import logging
import mimetypes
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass, field
import pandas as pd

logger = logging.getLogger(__name__)


@dataclass
class DataSchema:
    """Represents inferred data schema"""
    columns: Dict[str, str]  # column_name -> data_type
    row_count: int
    has_header: bool
    detected_patterns: List[str] = field(default_factory=list)


@dataclass
class StructuredDocument:
    """Represents a parsed document with structured data"""
    raw_text: str = ""
    tables: List[pd.DataFrame] = field(default_factory=list)
    images: List[Any] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    extracted_data: Dict[str, Any] = field(default_factory=dict)
    file_path: str = ""
    file_type: str = ""
    file_format: str = ""


class UniversalDocumentParser:
    """Parse ANY document format into structured data"""
    
    SUPPORTED_FORMATS = {
        'images': ['.jpg', '.png', '.jpeg', '.heic', '.bmp', '.tiff', '.gif'],
        'documents': ['.pdf', '.docx', '.doc', '.txt', '.rtf'],
        'spreadsheets': ['.xlsx', '.xls', '.csv', '.ods', '.tsv'],
        'data': ['.json', '.xml', '.yaml', '.yml', '.html']
    }
    
    def __init__(self):
        """Initialize the document parser"""
        self.ocr_engine = None  # Will be initialized when needed
        self.pdf_parser = None  # Will be initialized when needed
        self.image_processor = None  # Will be initialized when needed
        logger.info("UniversalDocumentParser initialized")
    
    def detect_format(self, file_path: str) -> Dict[str, Any]:
        """
        Auto-detect file format and type
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with:
                - type: 'image'|'document'|'spreadsheet'|'data'
                - format: file extension (e.g., '.pdf')
                - mime: MIME type
                - size: file size in bytes
        """
        path = Path(file_path)
        
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return {
                'type': 'unknown',
                'format': '',
                'mime': '',
                'size': 0,
                'exists': False
            }
        
        # Get file extension
        file_format = path.suffix.lower()
        
        # Get MIME type
        mime_type, _ = mimetypes.guess_type(str(path))
        
        # Get file size
        file_size = path.stat().st_size
        
        # Determine file type category
        file_type = 'unknown'
        for category, extensions in self.SUPPORTED_FORMATS.items():
            if file_format in extensions:
                file_type = category
                break
        
        result = {
            'type': file_type,
            'format': file_format,
            'mime': mime_type or 'application/octet-stream',
            'size': file_size,
            'exists': True,
            'name': path.name
        }
        
        logger.info(f"Detected file: {path.name} ({file_type}, {file_format}, {file_size} bytes)")
        
        return result
    
    def parse(self, file_path: str, options: Optional[Dict] = None) -> StructuredDocument:
        """
        Universal parser - automatically routes to appropriate handler
        
        Args:
            file_path: Path to file to parse
            options: Optional parsing options
            
        Returns:
            StructuredDocument with parsed content
        """
        if options is None:
            options = {}
        
        # Detect file format
        file_info = self.detect_format(file_path)
        
        if not file_info['exists']:
            logger.error(f"Cannot parse non-existent file: {file_path}")
            return StructuredDocument(
                file_path=file_path,
                file_type='unknown',
                file_format='',
                metadata={'error': 'File not found'}
            )
        
        file_type = file_info['type']
        file_format = file_info['format']
        
        # Route to appropriate parser
        try:
            if file_type == 'spreadsheets':
                return self._parse_spreadsheet(file_path, file_format, file_info)
            elif file_type == 'documents':
                return self._parse_document(file_path, file_format, file_info)
            elif file_type == 'images':
                return self._parse_image(file_path, file_format, file_info)
            elif file_type == 'data':
                return self._parse_data_file(file_path, file_format, file_info)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return self._parse_as_text(file_path, file_info)
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {e}")
            return StructuredDocument(
                file_path=file_path,
                file_type=file_type,
                file_format=file_format,
                metadata={'error': str(e)}
            )
    
    def _parse_spreadsheet(self, file_path: str, file_format: str, 
                          file_info: Dict) -> StructuredDocument:
        """Parse spreadsheet files (Excel, CSV, etc.)"""
        tables = []
        
        try:
            if file_format in ['.xlsx', '.xls']:
                # Parse Excel file
                df = pd.read_excel(file_path)
                tables.append(df)
            elif file_format == '.csv':
                # Parse CSV file
                df = pd.read_csv(file_path)
                tables.append(df)
            elif file_format == '.tsv':
                # Parse TSV file
                df = pd.read_csv(file_path, sep='\t')
                tables.append(df)
            
            logger.info(f"Parsed spreadsheet: {len(tables)} table(s) found")
            
            return StructuredDocument(
                tables=tables,
                file_path=file_path,
                file_type='spreadsheets',
                file_format=file_format,
                metadata={
                    **file_info,
                    'table_count': len(tables),
                    'row_count': tables[0].shape[0] if tables else 0,
                    'column_count': tables[0].shape[1] if tables else 0
                }
            )
        except Exception as e:
            logger.error(f"Error parsing spreadsheet {file_path}: {e}")
            return StructuredDocument(
                file_path=file_path,
                file_type='spreadsheets',
                file_format=file_format,
                metadata={**file_info, 'error': str(e)}
            )
    
    def _parse_document(self, file_path: str, file_format: str, 
                       file_info: Dict) -> StructuredDocument:
        """Parse document files (PDF, Word, etc.)"""
        raw_text = ""
        tables = []
        
        try:
            if file_format == '.txt':
                # Read plain text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    raw_text = f.read()
            elif file_format == '.pdf':
                # PDF parsing would require pdfplumber (not installed by default)
                logger.warning("PDF parsing requires pdfplumber - treating as placeholder")
                raw_text = f"[PDF Document: {Path(file_path).name}]"
            elif file_format in ['.docx', '.doc']:
                # Word document parsing would require python-docx (already in requirements)
                try:
                    from docx import Document
                    doc = Document(file_path)
                    raw_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                except ImportError:
                    logger.warning("python-docx not available - treating as placeholder")
                    raw_text = f"[Word Document: {Path(file_path).name}]"
            
            logger.info(f"Parsed document: {len(raw_text)} characters")
            
            return StructuredDocument(
                raw_text=raw_text,
                tables=tables,
                file_path=file_path,
                file_type='documents',
                file_format=file_format,
                metadata={
                    **file_info,
                    'text_length': len(raw_text)
                }
            )
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {e}")
            return StructuredDocument(
                file_path=file_path,
                file_type='documents',
                file_format=file_format,
                metadata={**file_info, 'error': str(e)}
            )
    
    def _parse_image(self, file_path: str, file_format: str, 
                    file_info: Dict) -> StructuredDocument:
        """Parse image files (requires OCR for text extraction)"""
        # Image parsing would require pytesseract and PIL
        # For now, create placeholder
        logger.info(f"Image file detected: {file_path}")
        logger.warning("OCR functionality requires pytesseract - treating as placeholder")
        
        return StructuredDocument(
            raw_text=f"[Image file: {Path(file_path).name}]",
            file_path=file_path,
            file_type='images',
            file_format=file_format,
            metadata={
                **file_info,
                'ocr_available': False
            }
        )
    
    def _parse_data_file(self, file_path: str, file_format: str, 
                        file_info: Dict) -> StructuredDocument:
        """Parse data files (JSON, XML, etc.)"""
        extracted_data = {}
        raw_text = ""
        
        try:
            if file_format == '.json':
                import json
                with open(file_path, 'r', encoding='utf-8') as f:
                    extracted_data = json.load(f)
                    raw_text = json.dumps(extracted_data, indent=2)
            elif file_format in ['.yaml', '.yml']:
                # YAML parsing would require PyYAML
                logger.warning("YAML parsing requires PyYAML - treating as text")
                with open(file_path, 'r', encoding='utf-8') as f:
                    raw_text = f.read()
            
            logger.info(f"Parsed data file: {file_format}")
            
            return StructuredDocument(
                raw_text=raw_text,
                extracted_data=extracted_data,
                file_path=file_path,
                file_type='data',
                file_format=file_format,
                metadata={
                    **file_info,
                    'has_structured_data': bool(extracted_data)
                }
            )
        except Exception as e:
            logger.error(f"Error parsing data file {file_path}: {e}")
            return StructuredDocument(
                file_path=file_path,
                file_type='data',
                file_format=file_format,
                metadata={**file_info, 'error': str(e)}
            )
    
    def _parse_as_text(self, file_path: str, file_info: Dict) -> StructuredDocument:
        """Fallback: try to parse as text"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                raw_text = f.read()
            
            return StructuredDocument(
                raw_text=raw_text,
                file_path=file_path,
                file_type='unknown',
                file_format=file_info.get('format', ''),
                metadata={
                    **file_info,
                    'parsed_as': 'text',
                    'text_length': len(raw_text)
                }
            )
        except Exception as e:
            logger.error(f"Error parsing as text {file_path}: {e}")
            return StructuredDocument(
                file_path=file_path,
                file_type='unknown',
                file_format=file_info.get('format', ''),
                metadata={**file_info, 'error': str(e)}
            )
    
    def infer_schema(self, data: pd.DataFrame) -> DataSchema:
        """
        Automatically infer data structure and types
        
        Args:
            data: DataFrame to analyze
            
        Returns:
            DataSchema with inferred information
        """
        columns = {}
        detected_patterns = []
        
        for col in data.columns:
            dtype = str(data[col].dtype)
            columns[col] = dtype
            
            # Detect common patterns
            if 'email' in col.lower():
                detected_patterns.append(f"{col}: email_field")
            elif 'name' in col.lower():
                detected_patterns.append(f"{col}: name_field")
            elif 'score' in col.lower() or 'result' in col.lower():
                detected_patterns.append(f"{col}: score_field")
        
        return DataSchema(
            columns=columns,
            row_count=len(data),
            has_header=True,
            detected_patterns=detected_patterns
        )
